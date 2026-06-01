import os
import json
import PyPDF2
try:
    import docx
except ImportError:
    docx = None

def read_txt(file_path):
    """Reads a text or markdown file with UTF-8, falling back to ISO-8859-1."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()

def read_pdf(file_path):
    """Reads a PDF file page by page and returns a list of pages with text."""
    pages = []
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append({
                    "content": text,
                    "metadata": {"page": idx + 1}
                })
    return pages

def read_docx(file_path):
    """Reads a DOCX file and returns the combined text of paragraphs and tables."""
    if docx is None:
        raise ImportError("python-docx is not installed or available.")
    
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

def format_json_value(val, indent=0):
    """Recursively formats a JSON value into a readable string."""
    prefix = "  " * indent
    if isinstance(val, dict):
        lines = []
        for k, v in val.items():
            formatted_v = format_json_value(v, indent + 1)
            if "\n" in formatted_v:
                lines.append(f"{prefix}{k}:\n{formatted_v}")
            else:
                lines.append(f"{prefix}{k}: {formatted_v.strip()}")
        return "\n".join(lines)
    elif isinstance(val, list):
        if all(isinstance(x, (str, int, float, bool)) for x in val):
            # Short list of primitive values
            return ", ".join(str(x) for x in val)
        else:
            lines = []
            for item in val:
                lines.append(format_json_value(item, indent + 1))
            return "\n".join(lines)
    else:
        return str(val)

def read_json(file_path):
    """Reads a JSON file and formats its content semantically.
    
    If the JSON contains a list of items (e.g. at the root, or inside a common list key),
    it yields each item as a separate chunk with root-level metadata merged.
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    chunks = []
    
    # Check if data is a list
    if isinstance(data, list):
        for idx, item in enumerate(data):
            formatted_text = format_json_value(item)
            chunks.append({
                "content": formatted_text,
                "metadata": {"item_index": idx}
            })
    # Check if there is a main content list in the dict (like 'contenido', 'data', 'items')
    elif isinstance(data, dict):
        list_keys = [k for k, v in data.items() if isinstance(v, list) and k in ('contenido', 'data', 'items', 'errors', 'errores')]
        
        # Extract base metadata from the dict (excluding lists)
        base_meta = {k: v for k, v in data.items() if not isinstance(v, (dict, list))}
        
        if list_keys:
            # We assume the first matching list key contains the main items
            main_key = list_keys[0]
            for idx, item in enumerate(data[main_key]):
                # Merge base metadata directly into the item dict to maintain single-paragraph coherence
                merged_item = {}
                for k, v in base_meta.items():
                    merged_item[f"doc_{k}"] = v
                
                if isinstance(item, dict):
                    for k, v in item.items():
                        merged_item[k] = v
                else:
                    merged_item["value"] = item
                    
                formatted_item = format_json_value(merged_item)
                
                # Merge base metadata into dictionary
                metadata = base_meta.copy()
                metadata["item_index"] = idx
                if isinstance(item, dict) and "id" in item and isinstance(item["id"], (str, int)):
                    metadata["item_id"] = str(item["id"])
                
                chunks.append({
                    "content": formatted_item,
                    "metadata": metadata
                })
        else:
            # No list found, treat the whole dict as one doc
            chunks.append({
                "content": format_json_value(data),
                "metadata": base_meta
            })
    else:
        chunks.append({
            "content": str(data),
            "metadata": {}
        })
        
    return chunks

def read_document(file_path):
    """Reads any supported document file.
    
    Returns a list of dicts: [{"content": str, "metadata": dict}]
    
    Handles errors gracefully, printing the error and raising it or returning empty.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext in ('.txt', '.md'):
            text = read_txt(file_path)
            return [{"content": text, "metadata": {}}]
        elif ext == '.pdf':
            return read_pdf(file_path)
        elif ext == '.docx':
            text = read_docx(file_path)
            return [{"content": text, "metadata": {}}]
        elif ext == '.json':
            return read_json(file_path)
        else:
            print(f"[Warning] Formato de archivo no soportado: {ext} para el archivo {file_path}")
            return []
    except Exception as e:
        print(f"[Error] No se pudo leer el archivo '{file_path}': {e}")
        return []

def read_all_documents(docs_dir):
    """Scans the given directory and reads all supported files.
    
    Returns a list of dicts: [{"content": str, "metadata": dict, "source": str, "filename": str}]
    """
    documents = []
    if not os.path.exists(docs_dir):
        print(f"[Error] La carpeta '{docs_dir}' no existe.")
        return documents
        
    for root, _, files in os.walk(docs_dir):
        for file in files:
            file_path = os.path.join(root, file)
            # Skip hidden files
            if file.startswith('.'):
                continue
                
            print(f"Leyendo: {file_path}...")
            file_docs = read_document(file_path)
            
            # Enrich metadata with file information
            for doc in file_docs:
                meta = doc.get("metadata", {})
                meta["source"] = file_path
                meta["filename"] = file
                
                documents.append({
                    "content": doc["content"],
                    "metadata": meta
                })
                
    return documents
